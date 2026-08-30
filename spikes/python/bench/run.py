"""Ferret EPIC-005 Python evaluation benchmarks.

Mirrors spikes/typescript/bench/run.mjs: same corpus, same iteration counts,
same result contract, so the two runtimes are compared on like work.
"""
import json
import os
import subprocess
import sys
import time
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from lib import (CODE, DOCS, LARGE, MALFORMED, REPO, emit, sha256,  # noqa: E402
                 stats, timed, walk)

HERE = Path(__file__).resolve().parent
B = {}


def bench(name):
    def deco(fn):
        B[name] = fn
        return fn
    return deco


# ---- memory: resident set after the core dependency surface is loaded ----
@bench("memory")
def _memory():
    import psutil
    proc = psutil.Process()
    before = proc.memory_info().rss
    import csv  # noqa: F401
    import docx  # noqa: F401
    import mcp.server.mcpserver  # noqa: F401
    import openpyxl  # noqa: F401
    import psycopg  # noqa: F401
    import pypdf  # noqa: F401
    import tree_sitter  # noqa: F401
    import tree_sitter_python  # noqa: F401
    after = proc.memory_info().rss
    emit("memory", "bytes", [after], {"rss_before_imports": before})


# ---- filesystem scan: walk + read + content hash ----
@bench("fsscan")
def _fsscan():
    def read_hash(f):
        data = f.read_bytes()
        sha256(data)
        return len(data)

    def sequential():
        files = walk(CODE)
        return len(files), sum(read_hash(f) for f in files)

    def threaded64():
        files = walk(CODE)
        with ThreadPoolExecutor(max_workers=64) as ex:
            return len(files), sum(ex.map(read_hash, files))

    # Both strategies are measured; each runtime is credited with its best.
    detail, meta = {}, {}
    for name, fn in (("sequential", sequential), ("threaded64", threaded64)):
        s = []
        for _ in range(5):
            with timed() as t:
                n, total = fn()
            meta = {"files": n, "bytes": total}
            s.append(t.ms)
        detail[name] = {"samples": s, **stats(s)}
    best = min(detail.items(), key=lambda kv: kv[1]["median"])
    emit("fsscan", "ms", detail[best[0]]["samples"],
         {**meta, "best_strategy": best[0], "concurrency": 64, "strategies": detail})


# ---- concurrent indexing: 4-process pool hashing the same tree ----
@bench("concurrency")
def _concurrency():
    from worker import hash_chunk
    files = [str(p) for p in walk(CODE)]
    n = 4
    chunks = [files[i::n] for i in range(n)]

    def processes():
        with ProcessPoolExecutor(max_workers=n) as ex:
            list(ex.map(hash_chunk, chunks))

    def threads():
        # hashlib and file I/O release the GIL, so threads are a real
        # alternative to paying Windows process-spawn cost.
        with ThreadPoolExecutor(max_workers=n) as ex:
            list(ex.map(hash_chunk, chunks))

    detail = {}
    for name, fn in (("ProcessPoolExecutor", processes), ("ThreadPoolExecutor", threads)):
        s = []
        for _ in range(3):
            with timed() as t:
                fn()
            s.append(t.ms)
        detail[name] = {"samples": s, **stats(s)}
    best = min(detail.items(), key=lambda kv: kv[1]["median"])
    emit("concurrency", "ms", detail[best[0]]["samples"],
         {"files": len(files), "workers": n, "best_model": best[0], "models": detail})


# ---- git: discovery/ingestion operations ----
@bench("git")
def _git():
    samples, meta = [], {}
    for _ in range(10):
        with timed() as t:
            # Run the three git calls concurrently, matching the Node
            # implementation's Promise.all rather than serialising them.
            cmds = [
                ["git", "-C", str(REPO), "log", "--format=%H%x00%an%x00%aI%x00%s", "-n", "200"],
                ["git", "-C", str(REPO), "ls-files"],
                ["git", "-C", str(REPO), "status", "--porcelain"],
            ]
            with ThreadPoolExecutor(max_workers=3) as ex:
                log, files, _status = list(ex.map(
                    lambda c: subprocess.run(c, capture_output=True, text=True, check=True), cmds))
            meta = {
                "commits": len([x for x in log.stdout.strip().split("\n") if x]),
                "tracked": len([x for x in files.stdout.strip().split("\n") if x]),
            }
        samples.append(t.ms)
    emit("git", "ms", samples, {**meta, "strategy": "git executable via subprocess"})


def parse_all(kind, ext, fn, runs=3):
    files = [f for f in walk(DOCS) if f.suffix == ext]
    samples, units = [], 0
    for _ in range(runs):
        with timed() as t:
            units = sum(fn(f) for f in files)
        samples.append(t.ms)
    emit(f"parse_{kind}", "ms", samples, {"files": len(files), "units": units})


@bench("pdf")
def _pdf():
    from pypdf import PdfReader

    def one(p):
        r = PdfReader(str(p))
        # Count extracted text runs per page to mirror the pdfjs item count.
        return sum(len((pg.extract_text() or "").split("\n")) for pg in r.pages)
    parse_all("pdf", ".pdf", one)


@bench("docx")
def _docx():
    import docx as pydocx

    def one(p):
        d = pydocx.Document(str(p))
        n = sum(len(par.text) for par in d.paragraphs)
        for tbl in d.tables:
            for row in tbl.rows:
                n += sum(len(c.text) for c in row.cells)
        return n
    parse_all("docx", ".docx", one)


@bench("xlsx")
def _xlsx():
    from openpyxl import load_workbook

    def one(p):
        wb = load_workbook(str(p))
        n = 0
        for ws in wb.worksheets:
            for _ in ws.iter_rows():
                n += 1
        return n
    parse_all("xlsx", ".xlsx", one)


@bench("csv")
def _csv():
    import csv as pycsv

    def one(p):
        n = 0
        with open(p, newline="", encoding="utf-8") as fh:
            for _ in pycsv.reader(fh):
                n += 1
        return n
    parse_all("csv", ".csv", one)


# ---- tree-sitter: symbol extraction over the code tree ----
@bench("treesitter")
def _treesitter():
    import tree_sitter_javascript
    import tree_sitter_python
    import tree_sitter_typescript
    from tree_sitter import Language, Parser

    langs = {
        ".ts": Language(tree_sitter_typescript.language_typescript()),
        ".js": Language(tree_sitter_javascript.language()),
        ".py": Language(tree_sitter_python.language()),
    }
    files = [f for f in walk(CODE) if f.suffix in langs]
    samples, meta = [], {}
    for _ in range(3):
        with timed() as t:
            nodes = parsed = 0
            for f in files:
                parser = Parser(langs[f.suffix])
                tree = parser.parse(f.read_bytes())
                stack = [tree.root_node]
                while stack:
                    node = stack.pop()
                    if node.is_named:
                        nodes += 1
                    stack.extend(node.named_children)
                parsed += 1
            meta = {"parsed": parsed, "named_nodes": nodes}
        samples.append(t.ms)
    emit("treesitter", "ms", samples, {**meta, "binding": "tree-sitter (native wheel)"})


# ---- large single file: streaming behaviour and peak resident memory ----
@bench("largefile")
def _largefile():
    import csv as pycsv
    import threading

    import psutil
    proc = psutil.Process()
    p = LARGE / "large.csv"
    size = p.stat().st_size
    samples, rows, peak = [], 0, 0
    for _ in range(3):
        stop = threading.Event()

        def sample_rss():
            nonlocal peak
            while not stop.is_set():
                peak = max(peak, proc.memory_info().rss)
                time.sleep(0.05)

        th = threading.Thread(target=sample_rss, daemon=True)
        th.start()
        with timed() as t:
            n = 0
            with open(p, newline="", encoding="utf-8") as fh:
                for _ in pycsv.reader(fh):
                    n += 1
            rows = n
        stop.set()
        th.join()
        samples.append(t.ms)
    emit("largefile_csv", "ms", samples,
         {"bytes": size, "rows": rows, "peak_rss_bytes": peak, "mode": "streaming"})


# ---- robustness: hostile/corrupt input, one isolated process per case ----
@bench("robustness")
def _robustness():
    manifest = json.loads((MALFORMED / "MALFORMED.json").read_text(encoding="utf-8"))
    timeout = 30
    results = []
    for c in manifest["malformed"]:
        f = MALFORMED / c["file"]
        t0 = time.perf_counter()
        try:
            cp = subprocess.run([sys.executable, str(HERE / "parse_one.py"), c["kind"], str(f)],
                                capture_output=True, text=True, timeout=timeout)
            out = cp.stdout.strip()
            if out:
                try:
                    res = json.loads(out)
                except json.JSONDecodeError:
                    res = {"outcome": "crash", "exit_code": cp.returncode,
                           "stderr": cp.stderr[-200:]}
            else:
                res = {"outcome": "crash", "exit_code": cp.returncode,
                       "stderr": cp.stderr[-200:]}
        except subprocess.TimeoutExpired:
            res = {"outcome": "timeout"}
        results.append({**c, **res, "ms": (time.perf_counter() - t0) * 1000.0})
    tally = {}
    for r in results:
        tally[r["outcome"]] = tally.get(r["outcome"], 0) + 1
    emit("robustness", "cases", [len(results)], {"tally": tally, "results": results})


# ---- postgres: bulk ingest, FTS index build, FTS query ----
@bench("postgres")
def _postgres():
    import psycopg
    rows_n = 50000
    with psycopg.connect(os.environ["FERRET_PG_URL"]) as conn:
        with conn.cursor() as cur:
            cur.execute("DROP TABLE IF EXISTS bench_py")
            cur.execute("CREATE TABLE bench_py (id int primary key, name text, "
                        "body text, tsv tsvector)")
            conn.commit()
            data = [(i, f"name{i}",
                     f"ferret evidence provenance row {i} indexed repository content")
                    for i in range(rows_n)]
            with timed() as ins:
                ch = 1000
                for i in range(0, len(data), ch):
                    chunk = data[i:i + ch]
                    args = []
                    for r in chunk:
                        args.extend(r)
                    ph = ",".join(["(%s,%s,%s)"] * len(chunk))
                    cur.execute(f"INSERT INTO bench_py (id,name,body) VALUES {ph}", args)
                conn.commit()
            with timed() as idx:
                cur.execute("UPDATE bench_py SET tsv = to_tsvector('english', body)")
                cur.execute("CREATE INDEX bench_py_tsv ON bench_py USING gin(tsv)")
                conn.commit()
            qs, hits = [], 0
            for _ in range(20):
                with timed() as q:
                    cur.execute("SELECT id FROM bench_py WHERE tsv @@ "
                                "plainto_tsquery('english','provenance evidence') LIMIT 100")
                    hits = len(cur.fetchall())
                qs.append(q.ms)
    emit("postgres_insert", "ms", [ins.ms], {"rows": rows_n, "driver": "psycopg3"})
    emit("postgres_index", "ms", [idx.ms], {"rows": rows_n, "driver": "psycopg3"})
    emit("postgres_fts_query", "ms", qs, {"hits": hits, "driver": "psycopg3"})


if __name__ == "__main__":
    which = sys.argv[1] if len(sys.argv) > 1 else ""
    if which not in B:
        print(f"unknown benchmark: {which}; have: {','.join(B)}", file=sys.stderr)
        sys.exit(2)
    B[which]()
