"""Parse a single file in an isolated process.

Hostile input must be observable (hang, crash, clean rejection) rather than
fatal to the whole benchmark run, so each case runs in its own process.
"""
import json
import sys


def parse_pdf(p):
    from pypdf import PdfReader
    r = PdfReader(p)
    return sum(len(pg.extract_text() or "") for pg in r.pages)


def parse_docx(p):
    import docx
    d = docx.Document(p)
    n = sum(len(par.text) for par in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            n += sum(len(c.text) for c in row.cells)
    return n


def parse_xlsx(p):
    from openpyxl import load_workbook
    wb = load_workbook(p)
    n = 0
    for ws in wb.worksheets:
        for _ in ws.iter_rows():
            n += 1
    return n


def parse_csv(p):
    import csv
    csv.field_size_limit(10_000_000)
    n = 0
    with open(p, newline="", encoding="utf-8", errors="replace") as fh:
        for _ in csv.reader(fh):
            n += 1
    return n


TABLE = {"pdf": parse_pdf, "docx": parse_docx, "xlsx": parse_xlsx, "csv": parse_csv}

if __name__ == "__main__":
    kind, file = sys.argv[1], sys.argv[2]
    try:
        units = TABLE[kind](file)
        sys.stdout.write(json.dumps({"outcome": "ok", "units": units}))
    except BaseException as e:  # noqa: BLE001 - classifying, not handling
        sys.stdout.write(json.dumps({
            "outcome": "clean_error",
            "error": type(e).__name__,
            "message": str(e)[:200],
        }))
