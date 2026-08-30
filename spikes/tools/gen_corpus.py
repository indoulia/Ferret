"""Generate the deterministic evaluation corpus shared by both spikes.

Generation is NOT a measured benchmark. Both runtimes parse the identical
bytes produced here, so any parser difference is attributable to the parser
and runtime rather than to differing inputs.

Usage: python spikes/tools/gen_corpus.py [--out DIR]
"""
import argparse
import csv
import json
import random
import shutil
import sys
from pathlib import Path

SEED = 20260830
CODE_FILES = 2000
PDF_COUNT, PDF_PAGES = 5, 20
DOCX_COUNT, DOCX_PARAS = 5, 400
XLSX_COUNT, XLSX_SHEETS, XLSX_ROWS = 5, 4, 2000
CSV_COUNT, CSV_ROWS = 5, 50_000

WORDS = ("ferret index evidence provenance repository worktree canonical entity "
         "relationship temporal retrieval hybrid provider contract migration "
         "checkpoint session parser symbol reference commit branch").split()


def sentence(rnd, n=12):
    return " ".join(rnd.choice(WORDS) for _ in range(n)).capitalize() + "."


def ts_file(rnd, i):
    lines = [f"// generated module {i}", "import { strict as assert } from 'node:assert';", ""]
    for f in range(rnd.randint(4, 14)):
        lines += [f"export interface Shape{i}_{f} {{ id: string; count: number; }}",
                  f"export function handle{i}_{f}(input: Shape{i}_{f}): number {{",
                  f"  // {sentence(rnd, 8)}", "  let total = 0;",
                  "  for (let k = 0; k < input.count; k++) { total += k; }",
                  "  assert.ok(total >= 0);", "  return total;", "}", ""]
    return "\n".join(lines)


def py_file(rnd, i):
    lines = [f'"""generated module {i}."""', "from dataclasses import dataclass", ""]
    for f in range(rnd.randint(4, 14)):
        lines += ["@dataclass", f"class Shape{i}_{f}:", "    id: str", "    count: int", "",
                  f"def handle_{i}_{f}(value: Shape{i}_{f}) -> int:", f'    """{sentence(rnd, 8)}"""',
                  "    total = 0", "    for k in range(value.count):", "        total += k",
                  "    return total", ""]
    return "\n".join(lines)


def js_file(rnd, i):
    lines = [f"// generated module {i}", "'use strict';", ""]
    for f in range(rnd.randint(4, 12)):
        lines += [f"function handle{i}_{f}(input) {{", f"  // {sentence(rnd, 8)}",
                  "  let total = 0;", "  for (let k = 0; k < input.count; k++) total += k;",
                  "  return total;", "}", f"module.exports.handle{i}_{f} = handle{i}_{f};", ""]
    return "\n".join(lines)


def md_file(rnd, i):
    lines = [f"# Document {i}", ""]
    for s in range(rnd.randint(3, 9)):
        lines += [f"## Section {s}", "", sentence(rnd, 20), "", "- " + sentence(rnd, 6),
                  "- " + sentence(rnd, 6), ""]
    return "\n".join(lines)


def json_file(rnd, i):
    return json.dumps({"id": i, "entries": [{"k": f"key{n}", "v": rnd.randint(0, 9999),
                                             "note": sentence(rnd, 5)} for n in range(rnd.randint(20, 90))]}, indent=2)


def gen_code_tree(out, rnd):
    mix = [("ts", ts_file, 0.40), ("py", py_file, 0.30), ("js", js_file, 0.10),
           ("md", md_file, 0.15), ("json", json_file, 0.05)]
    made = 0
    for i in range(CODE_FILES):
        r, acc = rnd.random(), 0.0
        ext, fn = mix[-1][0], mix[-1][1]
        for e, f, w in mix:
            acc += w
            if r <= acc:
                ext, fn = e, f
                break
        depth = rnd.randint(1, 5)
        d = out.joinpath("code", *[f"pkg{rnd.randint(0, 6)}" for _ in range(depth)])
        d.mkdir(parents=True, exist_ok=True)
        (d / f"mod_{i}.{ext}").write_text(fn(rnd, i), encoding="utf-8")
        made += 1
    return made


def gen_pdfs(out, rnd):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas
    d = out / "docs"
    d.mkdir(parents=True, exist_ok=True)
    for i in range(PDF_COUNT):
        c = canvas.Canvas(str(d / f"doc_{i}.pdf"), pagesize=LETTER)
        for p in range(PDF_PAGES):
            c.setFont("Helvetica-Bold", 14)
            c.drawString(72, 720, f"Document {i} - Page {p + 1}")
            c.setFont("Helvetica", 10)
            y = 690
            for _ in range(40):
                c.drawString(72, y, sentence(rnd, 11)[:95])
                y -= 15
            c.showPage()
        c.save()
    return PDF_COUNT


def gen_docx(out, rnd):
    from docx import Document
    d = out / "docs"
    d.mkdir(parents=True, exist_ok=True)
    for i in range(DOCX_COUNT):
        doc = Document()
        doc.add_heading(f"Specification {i}", level=1)
        for p in range(DOCX_PARAS):
            if p % 50 == 0:
                doc.add_heading(f"Section {p // 50}", level=2)
            doc.add_paragraph(sentence(rnd, 24))
        t = doc.add_table(rows=1, cols=4)
        for row in range(60):
            cells = t.add_row().cells
            for ci in range(4):
                cells[ci].text = f"r{row}c{ci}"
        doc.save(str(d / f"spec_{i}.docx"))
    return DOCX_COUNT


def gen_xlsx(out, rnd):
    from openpyxl import Workbook
    d = out / "docs"
    d.mkdir(parents=True, exist_ok=True)
    for i in range(XLSX_COUNT):
        wb = Workbook()
        wb.remove(wb.active)
        for s in range(XLSX_SHEETS):
            ws = wb.create_sheet(f"sheet{s}")
            ws.append(["id", "name", "amount", "note"])
            for r in range(XLSX_ROWS):
                ws.append([r, f"name{r}", rnd.randint(0, 100000) / 100.0, sentence(rnd, 6)])
        wb.save(str(d / f"book_{i}.xlsx"))
    return XLSX_COUNT


def gen_csv(out, rnd):
    d = out / "docs"
    d.mkdir(parents=True, exist_ok=True)
    for i in range(CSV_COUNT):
        with (d / f"table_{i}.csv").open("w", newline="", encoding="utf-8") as fh:
            w = csv.writer(fh)
            w.writerow(["id", "name", "amount", "note"])
            for r in range(CSV_ROWS):
                w.writerow([r, f"name{r}", rnd.randint(0, 100000) / 100.0, sentence(rnd, 6)])
    return CSV_COUNT


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default=str(Path(__file__).resolve().parents[1] / "corpus"))
    ap.add_argument("--force", action="store_true")
    args = ap.parse_args()
    out = Path(args.out)
    marker = out / "CORPUS.json"
    if marker.exists() and not args.force:
        print(json.dumps({"status": "exists", **json.loads(marker.read_text())}))
        return 0
    if out.exists():
        shutil.rmtree(out)
    out.mkdir(parents=True)
    rnd = random.Random(SEED)
    manifest = {
        "seed": SEED,
        "code_files": gen_code_tree(out, rnd),
        "pdf": gen_pdfs(out, rnd),
        "docx": gen_docx(out, rnd),
        "xlsx": gen_xlsx(out, rnd),
        "csv": gen_csv(out, rnd),
    }
    manifest["bytes"] = sum(p.stat().st_size for p in out.rglob("*") if p.is_file())
    marker.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps({"status": "generated", **manifest}))
    return 0


if __name__ == "__main__":
    sys.exit(main())
